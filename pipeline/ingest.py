"""Ingest stage: referral file -> Claude-ready content blocks.

txt  -> text block (read directly)
docx -> text block (paragraphs + tables via python-docx)
pdf  -> document block (Claude reads PDFs natively)
png/jpg -> image block
heic -> converted to jpeg in memory (Claude does not accept heic), then image block
"""

import base64
import io
from pathlib import Path

import pillow_heif
from docx import Document
from PIL import Image

pillow_heif.register_heif_opener()

SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf", ".png", ".jpg", ".jpeg", ".heic"}

# Anthropic rejects images over 10MB once base64-encoded (~4/3 of raw bytes).
MAX_RAW_IMAGE_BYTES = int(10_485_760 * 3 / 4 * 0.95)

_IMAGE_MEDIA_TYPES = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg"}


def _text_block(text: str) -> dict:
    return {"type": "text", "text": text}


def _base64_block(block_type: str, media_type: str, data: bytes) -> dict:
    return {
        "type": block_type,
        "source": {
            "type": "base64",
            "media_type": media_type,
            "data": base64.b64encode(data).decode(),
        },
    }


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _to_jpeg(image: Image.Image, quality: int = 90) -> bytes:
    buffer = io.BytesIO()
    image.convert("RGB").save(buffer, format="JPEG", quality=quality)
    return buffer.getvalue()


def _heic_to_jpeg(path: Path) -> bytes:
    return _to_jpeg(Image.open(path))


def _image_within_limit(path: Path) -> tuple[bytes, str]:
    """Return (bytes, media_type), recompressing oversized scans to fit the API cap."""
    data = path.read_bytes()
    if len(data) <= MAX_RAW_IMAGE_BYTES:
        return data, _IMAGE_MEDIA_TYPES[path.suffix.lower()]
    image = Image.open(path)
    for quality in (90, 80, 70):
        data = _to_jpeg(image, quality)
        if len(data) <= MAX_RAW_IMAGE_BYTES:
            return data, "image/jpeg"
    # Halve the resolution until the encoded size actually fits (still ample
    # for reading typed text); below ~200px there is nothing left to read.
    while image.width >= 400 and image.height >= 400:
        image = image.reduce(2)
        data = _to_jpeg(image, 80)
        if len(data) <= MAX_RAW_IMAGE_BYTES:
            return data, "image/jpeg"
    raise ValueError(
        f"{path.name} cannot be compressed under the API image size limit"
    )


def ingest(path: Path) -> list[dict]:
    """Return the Claude message content blocks for one referral file."""
    ext = path.suffix.lower()
    if ext == ".txt":
        return [_text_block(path.read_text())]
    if ext == ".docx":
        return [_text_block(_docx_text(path))]
    if ext == ".pdf":
        return [_base64_block("document", "application/pdf", path.read_bytes())]
    if ext in _IMAGE_MEDIA_TYPES:
        data, media_type = _image_within_limit(path)
        return [_base64_block("image", media_type, data)]
    if ext == ".heic":
        return [_base64_block("image", "image/jpeg", _heic_to_jpeg(path))]
    raise ValueError(f"Unsupported file type: {path.name}")
